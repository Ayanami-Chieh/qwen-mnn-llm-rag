"""
文档管理器模块 - 支持多种文件格式的加载和处理
支持格式：TXT, Markdown, DOCX, PDF
分块策略：递归字符分块 / 滑动窗口分块（纯 Python 实现，无需 LangChain）
"""
import os
import re
import logging
from pathlib import Path
from typing import List, Dict, Tuple, Optional
from abc import ABC, abstractmethod
import json
from datetime import datetime
from enum import Enum

logger = logging.getLogger(__name__)


# ============================================================================
# 分块策略枚举
# ============================================================================

class ChunkStrategy(Enum):
    """文本分块策略"""
    RECURSIVE = "recursive"        # 递归字符分块（纯 Python）
    SLIDING_WINDOW = "sliding_window"  # 滑动窗口分块（纯 Python）


# ============================================================================
# 文档加载器基类
# ============================================================================

class DocumentLoader(ABC):
    """文档加载器基类"""

    @abstractmethod
    def load(self, file_path: str) -> List[str]:
        """
        加载文档并返回文本内容

        Args:
            file_path: 文件路径

        Returns:
            文本列表，每个元素是一个段落
        """
        pass

    @staticmethod
    def get_file_info(file_path: str) -> Dict:
        """获取文件信息"""
        path = Path(file_path)
        return {
            'name': path.name,
            'size': path.stat().st_size,
            'created': datetime.fromtimestamp(path.stat().st_ctime).isoformat(),
            'modified': datetime.fromtimestamp(path.stat().st_mtime).isoformat(),
        }


# ============================================================================
# TXT 文档加载器
# ============================================================================

class TxtLoader(DocumentLoader):
    """文本文档加载器"""

    def load(self, file_path: str) -> List[str]:
        """加载TXT文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()

            # 按行分割，过滤空行
            paragraphs = [p.strip() for p in content.split('\n') if p.strip()]
            logger.info(f"✅ TXT文件加载成功: {len(paragraphs)} 段")
            return paragraphs

        except Exception as e:
            logger.error(f"❌ TXT文件加载失败: {e}")
            return []


# ============================================================================
# Markdown 文档加载器
# ============================================================================

class MarkdownLoader(DocumentLoader):
    """Markdown文档加载器"""

    def load(self, file_path: str) -> List[str]:
        """加载Markdown文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()

            paragraphs = self._parse_markdown(content)
            logger.info(f"✅ Markdown文件加载成功: {len(paragraphs)} 段")
            return paragraphs

        except Exception as e:
            logger.error(f"❌ Markdown文件加载失败: {e}")
            return []

    @staticmethod
    def _parse_markdown(content: str) -> List[str]:
        """解析Markdown内容"""
        # 移除Markdown特殊字符但保留内容
        content = re.sub(r'#+\s+', '', content)  # 移除标题符号
        content = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', content)  # 转换链接
        content = re.sub(r'\*\*([^*]+)\*\*', r'\1', content)  # 移除粗体
        content = re.sub(r'__([^_]+)__', r'\1', content)  # 移除粗体
        content = re.sub(r'\*([^*]+)\*', r'\1', content)  # 移除斜体
        content = re.sub(r'_([^_]+)_', r'\1', content)  # 移除斜体
        content = re.sub(r'`([^`]+)`', r'\1', content)  # 移除代码块标记
        content = re.sub(r'```[\s\S]*?```', '', content)  # 移除代码块

        # 按段落分割
        paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
        return paragraphs


# ============================================================================
# DOCX 文档加载器
# ============================================================================

class DocxLoader(DocumentLoader):
    """DOCX文档加载器"""

    def load(self, file_path: str) -> List[str]:
        """加载DOCX文件"""
        try:
            from docx import Document

            doc = Document(file_path)
            paragraphs = []

            # 提取段落
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text.strip())

            # 提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' '.join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        paragraphs.append(row_text)

            logger.info(f"✅ DOCX文件加载成功: {len(paragraphs)} 段")
            return paragraphs

        except ImportError:
            logger.error("❌ 缺少python-docx库，请运行: pip install python-docx")
            return []
        except Exception as e:
            logger.error(f"❌ DOCX文件加载失败: {e}")
            return []


# ============================================================================
# PDF 文档加载器
# ============================================================================

class PdfLoader(DocumentLoader):
    """PDF文档加载器"""

    def load(self, file_path: str) -> List[str]:
        """加载PDF文件"""
        try:
            import pypdf

            paragraphs = []

            with open(file_path, 'rb') as f:
                pdf_reader = pypdf.PdfReader(f)

                for page_num, page in enumerate(pdf_reader.pages, 1):
                    text = page.extract_text()

                    if text.strip():
                        # 按行分割，过滤空行
                        lines = [line.strip() for line in text.split('\n') if line.strip()]
                        paragraphs.extend(lines)

            logger.info(f"✅ PDF文件加载成功: {len(paragraphs)} 段")
            return paragraphs

        except ImportError:
            logger.error("❌ 缺少pypdf库，请运行: pip install pypdf")
            return []
        except Exception as e:
            logger.error(f"❌ PDF文件加载失败: {e}")
            return []


# ============================================================================
# 文档管理器
# ============================================================================

class DocumentManager:
    """文档管理器 - 支持多种格式的文档加载和处理"""

    # 支持的文件格式
    SUPPORTED_FORMATS = {
        '.txt': TxtLoader,
        '.md': MarkdownLoader,
        '.markdown': MarkdownLoader,
        '.docx': DocxLoader,
        '.doc': DocxLoader,
        '.pdf': PdfLoader,
    }

    def __init__(
        self,
        chunk_size: int = 200,
        overlap: int = 50,
        strategy: ChunkStrategy = ChunkStrategy.RECURSIVE,
    ):
        """
        初始化文档管理器

        Args:
            chunk_size: 分块大小（字符数）
            overlap:    分块重叠字符数（递归分块时为重叠量；
                        滑动窗口时为步长 = chunk_size - overlap）
            strategy:   分块策略，ChunkStrategy.RECURSIVE 或
                        ChunkStrategy.SLIDING_WINDOW
        """
        self.chunk_size = chunk_size
        self.overlap = overlap
        self.strategy = strategy
        self.documents = {}       # 存储已加载的文档
        self.file_metadata = {}   # 存储文件元数据

        logger.info(
            f"✅ DocumentManager 初始化完成 | 策略={strategy.value} | "
            f"chunk_size={chunk_size} | overlap={overlap}"
        )

    # ------------------------------------------------------------------
    # LangChain 分块器（内部工厂）
    # ------------------------------------------------------------------

    @staticmethod
    def _recursive_split(text: str, separators: List[str],
                         chunk_size: int, overlap: int) -> List[str]:
        """
        纯 Python 递归字符分块（等价于 LangChain RecursiveCharacterTextSplitter）。

        算法：
          1. 依次尝试 separators，找到第一个在 text 中出现的分隔符。
          2. 用该分隔符切割，对每段递归处理（传入剩余分隔符列表）。
          3. 将结果按 chunk_size / overlap 合并成最终块。
        """
        # 找第一个有效分隔符
        separator = ""
        new_separators: List[str] = []
        for i, sep in enumerate(separators):
            if sep == "" or sep in text:
                separator = sep
                new_separators = separators[i + 1:]
                break

        # 按分隔符切割
        if separator:
            splits = text.split(separator)
        else:
            splits = list(text)

        # 对每段递归或直接收集
        good: List[str] = []
        for s in splits:
            s = s.strip()
            if not s:
                continue
            if len(s) <= chunk_size:
                good.append(s)
            else:
                # 还需要进一步切割
                good.extend(
                    DocumentManager._recursive_split(
                        s, new_separators, chunk_size, overlap
                    )
                )

        # 合并小块，保留 overlap
        chunks: List[str] = []
        current = ""
        for piece in good:
            candidate = (current + separator + piece).strip() if current else piece
            if len(candidate) <= chunk_size:
                current = candidate
            else:
                if current:
                    chunks.append(current)
                # overlap：从已有块的尾部取 overlap 字符作为新块起点
                if overlap > 0 and current:
                    tail = current[-overlap:]
                    current = (tail + separator + piece).strip()
                else:
                    current = piece
        if current:
            chunks.append(current)

        return chunks if chunks else [text[:chunk_size]]

    @staticmethod
    def _sliding_split(text: str, chunk_size: int, overlap: int) -> List[str]:
        """
        纯 Python 滑动窗口分块（等价于 LangChain CharacterTextSplitter(separator="")）。

        步长 = chunk_size - overlap
        """
        if not text:
            return []
        step = max(1, chunk_size - overlap)
        chunks = []
        start = 0
        while start < len(text):
            chunks.append(text[start: start + chunk_size])
            start += step
        return chunks

    # ------------------------------------------------------------------
    # 公共分块入口
    # ------------------------------------------------------------------

    def chunk_text(
        self,
        text: str,
        chunk_size: Optional[int] = None,
        overlap: Optional[int] = None,
        strategy: Optional[ChunkStrategy] = None,
    ) -> List[str]:
        """
        将文本分块（统一入口，根据 strategy 调度到对应纯 Python 实现）。
        接口与原版完全兼容，不再依赖 LangChain。

        Args:
            text:       输入文本
            chunk_size: 分块大小，None 则使用实例默认值
            overlap:    重叠/步长，None 则使用实例默认值
            strategy:   分块策略，None 则使用实例默认策略

        Returns:
            分块列表
        """
        if not text:
            return []

        use_strategy = strategy or self.strategy
        use_size = chunk_size or self.chunk_size
        use_overlap = overlap if overlap is not None else self.overlap

        if use_strategy == ChunkStrategy.RECURSIVE:
            separators = [
                "\n\n", "\n",
                "。", "！", "？", ".", "!", "?",
                "；", ";", "，", ",", " ", "",
            ]
            chunks = self._recursive_split(text, separators, use_size, use_overlap)

        elif use_strategy == ChunkStrategy.SLIDING_WINDOW:
            chunks = self._sliding_split(text, use_size, use_overlap)

        else:
            logger.warning(f"未知分块策略 {use_strategy}，回退到递归分块")
            separators = ["\n\n", "\n", "。", "！", "？", ".", "!", "?",
                          "；", ";", "，", ",", " ", ""]
            chunks = self._recursive_split(text, separators, use_size, use_overlap)

        logger.debug(
            f"chunk_text | strategy={use_strategy.value} | "
            f"input_len={len(text)} | chunks={len(chunks)}"
        )
        return chunks

    # ------------------------------------------------------------------
    # 文档加载
    # ------------------------------------------------------------------

    def load_document(self, file_path: str) -> Tuple[bool, str, List[str]]:
        """
        加载单个文档

        Args:
            file_path: 文件路径

        Returns:
            (是否成功, 消息, 文本列表)
        """
        file_path = str(file_path)

        if not os.path.exists(file_path):
            msg = f"文件不存在: {file_path}"
            logger.error(f"❌ {msg}")
            return False, msg, []

        ext = Path(file_path).suffix.lower()

        if ext not in self.SUPPORTED_FORMATS:
            supported = ', '.join(self.SUPPORTED_FORMATS.keys())
            msg = f"不支持的文件格式: {ext}。支持的格式: {supported}"
            logger.error(f"❌ {msg}")
            return False, msg, []

        loader = self.SUPPORTED_FORMATS[ext]()
        paragraphs = loader.load(file_path)

        if not paragraphs:
            msg = f"文件加载失败或内容为空: {file_path}"
            logger.error(f"❌ {msg}")
            return False, msg, []

        doc_name = Path(file_path).stem
        self.documents[doc_name] = paragraphs
        self.file_metadata[doc_name] = {
            'path': file_path,
            'format': ext,
            'file_info': loader.get_file_info(file_path),
            'paragraph_count': len(paragraphs),
            'total_chars': sum(len(p) for p in paragraphs),
            'loaded_at': datetime.now().isoformat(),
        }

        msg = (
            f"文档加载成功: {len(paragraphs)} 段，"
            f"总计 {self.file_metadata[doc_name]['total_chars']} 字符"
        )
        logger.info(f"✅ {msg}")
        return True, msg, paragraphs

    def load_documents_from_directory(
        self, directory: str
    ) -> Dict[str, Tuple[bool, str, List[str]]]:
        """
        从目录加载所有支持的文档

        Args:
            directory: 目录路径

        Returns:
            {文档名: (是否成功, 消息, 文本列表)}
        """
        results = {}

        if not os.path.isdir(directory):
            logger.error(f"❌ 目录不存在: {directory}")
            return results

        for ext in self.SUPPORTED_FORMATS.keys():
            for file_path in Path(directory).glob(f'*{ext}'):
                success, msg, paragraphs = self.load_document(str(file_path))
                doc_name = file_path.stem
                results[doc_name] = (success, msg, paragraphs)

        logger.info(f"📁 目录加载完成: {len(results)} 个文件")
        return results

    # ------------------------------------------------------------------
    # 文档处理（分块 + 清理）
    # ------------------------------------------------------------------

    def process_documents(
        self,
        chunk_size: Optional[int] = None,
        strategy: Optional[ChunkStrategy] = None,
    ) -> List[str]:
        """
        处理所有已加载的文档，进行清理和分块。

        Args:
            chunk_size: 分块大小，None 则使用实例默认值
            strategy:   分块策略，None 则使用实例默认策略

        Returns:
            处理后的文本分块列表（已去重）
        """
        use_strategy = strategy or self.strategy
        all_chunks = []

        logger.info(
            f"🔄 开始处理文档 | 文档数={len(self.documents)} | "
            f"策略={use_strategy.value}"
        )

        for doc_name, paragraphs in self.documents.items():
            logger.info(f"  处理文档: {doc_name} ({len(paragraphs)} 段)")
            doc_chunks = []

            for para in paragraphs:
                cleaned = self._clean_text(para)
                if cleaned:
                    chunks = self.chunk_text(
                        cleaned,
                        chunk_size=chunk_size,
                        strategy=use_strategy,
                    )
                    doc_chunks.extend(chunks)

            logger.info(f"  ✅ {doc_name}: {len(doc_chunks)} 个分块")
            all_chunks.extend(doc_chunks)

        # 全局去重
        all_chunks = self._remove_duplicates(all_chunks)
        logger.info(f"✅ 文档处理完成: 生成 {len(all_chunks)} 个文本块")
        return all_chunks

    # ------------------------------------------------------------------
    # 工具方法
    # ------------------------------------------------------------------

    def _clean_text(self, text: str) -> str:
        """清理文本"""
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(
            r'[^\u4e00-\u9fa5a-zA-Z0-9\s，。！？；：、\'"\'""''（）【】《》]',
            '',
            text,
        )
        return text.strip()

    def _remove_duplicates(self, fragments: List[str]) -> List[str]:
        """移除重复片段"""
        seen = set()
        unique_fragments = []
        for frag in fragments:
            if frag not in seen:
                seen.add(frag)
                unique_fragments.append(frag)
        removed = len(fragments) - len(unique_fragments)
        logger.info(f"⚙️ 去重完成，共移除 {removed} 个重复片段")
        return unique_fragments

    # ------------------------------------------------------------------
    # 统计 / 元数据
    # ------------------------------------------------------------------

    def get_document_stats(self) -> Dict:
        """获取文档统计信息"""
        stats = {
            'total_documents': len(self.documents),
            'total_paragraphs': sum(len(p) for p in self.documents.values()),
            'total_characters': sum(
                self.file_metadata[doc_name]['total_chars']
                for doc_name in self.documents.keys()
            ),
            'chunk_strategy': self.strategy.value,
            'chunk_size': self.chunk_size,
            'overlap': self.overlap,
            'documents': {},
        }

        for doc_name, metadata in self.file_metadata.items():
            stats['documents'][doc_name] = {
                'format': metadata['format'],
                'paragraphs': metadata['paragraph_count'],
                'characters': metadata['total_chars'],
                'path': metadata['path'],
            }

        return stats

    def export_metadata(self, output_path: str):
        """导出文件元数据为JSON"""
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(self.file_metadata, f, ensure_ascii=False, indent=2)
            logger.info(f"✅ 元数据已导出: {output_path}")
        except Exception as e:
            logger.error(f"❌ 元数据导出失败: {e}")

    def clear_documents(self):
        """清空所有已加载的文档"""
        self.documents.clear()
        self.file_metadata.clear()
        logger.info("✅ 所有文档已清空")

    def get_loaded_documents_count(self) -> int:
        """获取已加载的文档数量"""
        return len(self.documents)

    # ------------------------------------------------------------------
    # 策略切换（运行时）
    # ------------------------------------------------------------------

    def set_strategy(
        self,
        strategy: ChunkStrategy,
        chunk_size: Optional[int] = None,
        overlap: Optional[int] = None,
    ):
        """
        运行时切换分块策略（同时可更新 chunk_size / overlap）。
        切换后会重置对应缓存的分块器，下次使用时重新创建。

        Args:
            strategy:   新的分块策略
            chunk_size: 新的分块大小（可选）
            overlap:    新的重叠大小（可选）
        """
        self.strategy = strategy
        if chunk_size is not None:
            self.chunk_size = chunk_size
        if overlap is not None:
            self.overlap = overlap

        logger.info(
            f"✅ 分块策略已切换 | strategy={strategy.value} | "
            f"chunk_size={self.chunk_size} | overlap={self.overlap}"
        )
